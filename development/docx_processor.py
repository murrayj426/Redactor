"""
Word Document Processor
Extract text from .docx files for audit system
"""

from docx import Document
import os

def extract_text_from_docx(file_path):
    """Extract text from a Word document"""
    try:
        doc = Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only include non-empty paragraphs
                full_text.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    
    except Exception as e:
        return f"Error extracting text from Word document: {str(e)}"

def save_docx_as_txt(docx_path, txt_path):
    """Convert Word document to text file"""
    text_content = extract_text_from_docx(docx_path)
    
    try:
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write("# Incident Handling and Documentation Standards (from Word document)\n\n")
            f.write(text_content)
        
        return f"Successfully converted {docx_path} to {txt_path}"
    
    except Exception as e:
        return f"Error saving text file: {str(e)}"

if __name__ == "__main__":
    docx_file = "Incident Handling and Documentation Standards.docx"
    txt_file = "incident_handling_procedures_full.txt"
    
    if os.path.exists(docx_file):
        result = save_docx_as_txt(docx_file, txt_file)
        print(result)
        
        # Also print a preview
        text_content = extract_text_from_docx(docx_file)
        print(f"\nDocument length: {len(text_content)} characters")
        print(f"Preview (first 500 characters):\n{text_content[:500]}...")
    else:
        print(f"Word document not found: {docx_file}")
